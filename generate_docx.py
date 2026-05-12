#!/usr/bin/env python3
"""Generate Word (.docx) - Data TFL Riska & Ujang BSPS 2026 Tahap IV
   Each desa on separate page. Green highlight for pengganti.
   Includes Tidak ACC section for LOA.
"""

import json, subprocess
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

ACCENT = RGBColor(0xBD, 0x26, 0x3F)
TEXT_PRIMARY = RGBColor(0x23, 0x22, 0x20)
TEXT_MUTED = RGBColor(0x78, 0x74, 0x6D)
RED_TEXT = RGBColor(0x72, 0x1C, 0x24)
GREEN_BG = "D4EDDA"
SURFACE_BG = "E7E4DF"
WHITE = "FFFFFF"
ACCENT_HEX = "BD263F"

output_path = "/home/z/my-project/upload/Data_Riska_Ujang_TFL_IBUN.docx"

# ━━ Read Data ━━
result = subprocess.run(['bun', '-e', '''
const XLSX = require("xlsx");
const wb = XLSX.readFile("upload/Data_Riska_Ujang_TFL_IBUN.xlsx");
const out = {};
const s1 = wb.Sheets["Data Penerima"];
const r1 = XLSX.utils.sheet_to_json(s1, {header:1, defval:""});
out.penerima = [];
for(let i=1; i<r1.length; i++){
  const r = r1[i];
  if(!r[1]) continue;
  out.penerima.push({
    no: r[0], nama: String(r[1]), lp: String(r[2]),
    ktp: String(r[3]), kk: String(r[4]), alamat: String(r[5]),
    desa: String(r[6]).trim(), kec: String(r[7]).trim(), ket: String(r[10]),
    pengganti: String(r[10]).includes("Pengganti") || String(r[10]).includes("Baru"),
  });
}
const s3 = wb.Sheets["Tidak ACC - LOA"];
out.tidak_acc = [];
if(s3){
  const r3 = XLSX.utils.sheet_to_json(s3, {header:1, defval:""});
  for(let i=1; i<r3.length; i++){
    const r = r3[i];
    if(!r[1]) continue;
    out.tidak_acc.push({
      no: r[0], nama: String(r[1]), lp: String(r[2]),
      ktp: String(r[3]), kk: String(r[4]), alamat: String(r[5]),
      alasan: String(r[6]),
    });
  }
}
console.log(JSON.stringify(out));
'''], capture_output=True, text=True, cwd='/home/z/my-project')
data = json.loads(result.stdout)
penerima = data['penerima']
tidak_acc = data['tidak_acc']

grouped = {}
for d in ['CIBEET', 'LAMPEGAN', 'LOA']:
    grouped[d] = [r for r in penerima if r.get('desa') == d]
total_l = sum(1 for r in penerima if r['lp'] == 'L')
total_p = sum(1 for r in penerima if r['lp'] == 'P')

# ━━ Helpers ━━
def shade(cell, color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def txt(cell, text, bold=False, color=None, size=8, align='center'):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if bold: run.bold = True
    if color: run.font.color.rgb = color

def borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '</w:tblBorders>'))
    if tbl.tblPr is None: tbl.insert(0, tblPr)

def add_page(doc, title_text, desa_name, rows, highlight_indices=None):
    """Add a full page with header + desa table."""
    if highlight_indices is None:
        highlight_indices = set()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Data Calon Penerima BSPS 2026 Tahap IV')
    run.bold = True; run.font.size = Pt(14); run.font.color.rgb = TEXT_PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TFL Teknik: Riska Arie Haryanti | TFL Pemberdayaan: Ujang Supriatna')
    run.font.size = Pt(10); run.font.color.rgb = TEXT_MUTED

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run(f'Desa {desa_name}')
    run.bold = True; run.font.size = Pt(11); run.font.color.rgb = ACCENT

    label = ' (baris hijau = data pengganti/baru)' if highlight_indices else ''
    p = doc.add_paragraph()
    run = p.add_run(f'Jumlah penerima ACC: {len(rows)} orang{label}')
    run.font.size = Pt(9); run.font.color.rgb = TEXT_MUTED

    doc.add_paragraph()

    headers = ['No', 'Nama', 'L/P', 'No KTP', 'No KK', 'Alamat', 'Desa', 'Kecamatan', 'Keterangan']
    widths = [0.6, 2.2, 0.5, 1.8, 1.8, 2.8, 1.0, 1.0, 1.8]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(table)

    for i, w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = Cm(w)

    for j, h in enumerate(headers):
        shade(table.rows[0].cells[j], ACCENT_HEX)
        txt(table.rows[0].cells[j], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=8)

    for i, r in enumerate(rows, 1):
        ket = '\u2605 Data Pengganti' if i in highlight_indices else (r.get('ket') or '-')
        vals = [r['no'], r['nama'], r['lp'], r['ktp'], r['kk'], r['alamat'], r['desa'], r['kec'], ket]
        for j, v in enumerate(vals):
            a = 'left' if j in [1, 5, 8] else 'center'
            if i in highlight_indices:
                shade(table.rows[i].cells[j], GREEN_BG)
            elif i % 2 == 0:
                shade(table.rows[i].cells[j], SURFACE_BG)
            else:
                shade(table.rows[i].cells[j], WHITE)
            txt(table.rows[i].cells[j], v, size=8, align=a)

# ━━ Build Document ━━
doc = Document()
section = doc.sections[0]
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)

# Title + summary
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Data Calon Penerima BSPS 2026 Tahap IV')
run.bold = True; run.font.size = Pt(14); run.font.color.rgb = TEXT_PRIMARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TFL Teknik: Riska Arie Haryanti | TFL Pemberdayaan: Ujang Supriatna')
run.font.size = Pt(10); run.font.color.rgb = TEXT_MUTED

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Kecamatan: IBUN, PASEH  |  Total ACC: {len(penerima)} orang  |  Laki-laki: {total_l}  |  Perempuan: {total_p}')
run.font.size = Pt(9); run.font.color.rgb = TEXT_MUTED

doc.add_paragraph()

# CIBEET
p = doc.add_paragraph()
run = p.add_run(f'Desa CIBEET - Kecamatan IBUN')
run.bold = True; run.font.size = Pt(11); run.font.color.rgb = ACCENT
p = doc.add_paragraph()
run = p.add_run(f'Jumlah penerima ACC: {len(grouped["CIBEET"])} orang')
run.font.size = Pt(9); run.font.color.rgb = TEXT_MUTED
doc.add_paragraph()
add_page.__code__  # reference to keep function available

# Build CIBEET table directly
headers = ['No', 'Nama', 'L/P', 'No KTP', 'No KK', 'Alamat', 'Desa', 'Kecamatan', 'Keterangan']
widths = [0.6, 2.2, 0.5, 1.8, 1.8, 2.8, 1.0, 1.0, 1.8]
table = doc.add_table(rows=1 + len(grouped['CIBEET']), cols=len(headers))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
borders(table)
for i, w in enumerate(widths):
    for row in table.rows: row.cells[i].width = Cm(w)
for j, h in enumerate(headers):
    shade(table.rows[0].cells[j], ACCENT_HEX)
    txt(table.rows[0].cells[j], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=8)
for i, r in enumerate(grouped['CIBEET'], 1):
    vals = [r['no'], r['nama'], r['lp'], r['ktp'], r['kk'], r['alamat'], r['desa'], r['kec'], 'ACC']
    for j, v in enumerate(vals):
        a = 'left' if j in [1, 5, 8] else 'center'
        shade(table.rows[i].cells[j], WHITE if i%2==1 else SURFACE_BG)
        txt(table.rows[i].cells[j], v, size=8, align=a)

# PAGE BREAK - LAMPEGAN
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Data Calon Penerima BSPS 2026 Tahap IV')
run.bold = True; run.font.size = Pt(14); run.font.color.rgb = TEXT_PRIMARY
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TFL Teknik: Riska Arie Haryanti | TFL Pemberdayaan: Ujang Supriatna')
run.font.size = Pt(10); run.font.color.rgb = TEXT_MUTED
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Desa LAMPEGAN - Kecamatan IBUN')
run.bold = True; run.font.size = Pt(11); run.font.color.rgb = ACCENT
p = doc.add_paragraph()
run = p.add_run(f'Jumlah penerima ACC: {len(grouped["LAMPEGAN"])} orang (baris hijau = data pengganti)')
run.font.size = Pt(9); run.font.color.rgb = TEXT_MUTED
doc.add_paragraph()

hl_lamp = {i+1 for i, r in enumerate(grouped['LAMPEGAN']) if r['pengganti']}
table = doc.add_table(rows=1 + len(grouped['LAMPEGAN']), cols=len(headers))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
borders(table)
for i, w in enumerate(widths):
    for row in table.rows: row.cells[i].width = Cm(w)
for j, h in enumerate(headers):
    shade(table.rows[0].cells[j], ACCENT_HEX)
    txt(table.rows[0].cells[j], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=8)
for i, r in enumerate(grouped['LAMPEGAN'], 1):
    ket = '\u2605 Data Pengganti' if i in hl_lamp else 'ACC'
    vals = [r['no'], r['nama'], r['lp'], r['ktp'], r['kk'], r['alamat'], r['desa'], r['kec'], ket]
    for j, v in enumerate(vals):
        a = 'left' if j in [1, 5, 8] else 'center'
        if i in hl_lamp: shade(table.rows[i].cells[j], GREEN_BG)
        elif i%2==0: shade(table.rows[i].cells[j], SURFACE_BG)
        else: shade(table.rows[i].cells[j], WHITE)
        txt(table.rows[i].cells[j], v, size=8, align=a)

# PAGE BREAK - LOA (ACC + Tidak ACC)
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Data Calon Penerima BSPS 2026 Tahap IV')
run.bold = True; run.font.size = Pt(14); run.font.color.rgb = TEXT_PRIMARY
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TFL Teknik: Riska Arie Haryanti | TFL Pemberdayaan: Ujang Supriatna')
run.font.size = Pt(10); run.font.color.rgb = TEXT_MUTED
doc.add_paragraph()

# LOA ACC section
p = doc.add_paragraph()
run = p.add_run('Desa LOA - Kecamatan PASEH')
run.bold = True; run.font.size = Pt(11); run.font.color.rgb = ACCENT
p = doc.add_paragraph()
run = p.add_run(f'Jumlah penerima ACC: {len(grouped["LOA"])} orang (baris hijau = data pengganti/baru)')
run.font.size = Pt(9); run.font.color.rgb = TEXT_MUTED
doc.add_paragraph()

hl_loa = {i+1 for i, r in enumerate(grouped['LOA']) if r['pengganti']}
table = doc.add_table(rows=1 + len(grouped['LOA']), cols=len(headers))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
borders(table)
for i, w in enumerate(widths):
    for row in table.rows: row.cells[i].width = Cm(w)
for j, h in enumerate(headers):
    shade(table.rows[0].cells[j], ACCENT_HEX)
    txt(table.rows[0].cells[j], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=8)
for i, r in enumerate(grouped['LOA'], 1):
    ket = '\u2605 Data Pengganti' if i in hl_loa else 'ACC'
    vals = [r['no'], r['nama'], r['lp'], r['ktp'], r['kk'], r['alamat'], r['desa'], r['kec'], ket]
    for j, v in enumerate(vals):
        a = 'left' if j in [1, 5, 8] else 'center'
        if i in hl_loa: shade(table.rows[i].cells[j], GREEN_BG)
        elif i%2==0: shade(table.rows[i].cells[j], SURFACE_BG)
        else: shade(table.rows[i].cells[j], WHITE)
        txt(table.rows[i].cells[j], v, size=8, align=a)

doc.add_paragraph()

# TIDAK ACC section
p = doc.add_paragraph()
run = p.add_run('Data Tidak ACC - Desa LOA')
run.bold = True; run.font.size = Pt(10); run.font.color.rgb = RED_TEXT
p = doc.add_paragraph()
run = p.add_run(f'Jumlah tidak ACC: {len(tidak_acc)} orang (11 Layak Huni, 1 Tidak Melanjutkan)')
run.font.size = Pt(9); run.font.color.rgb = TEXT_MUTED
doc.add_paragraph()

na_headers = ['No', 'Nama', 'L/P', 'No KTP', 'No KK', 'Alamat', 'Alasan Tidak ACC']
na_widths = [0.6, 2.2, 0.5, 1.8, 1.8, 2.8, 2.5]
table = doc.add_table(rows=1 + len(tidak_acc), cols=len(na_headers))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
borders(table)
for i, w in enumerate(na_widths):
    for row in table.rows: row.cells[i].width = Cm(w)
for j, h in enumerate(na_headers):
    shade(table.rows[0].cells[j], "6C757D")
    txt(table.rows[0].cells[j], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=8)
RED_BG = "FFF3CD"
for i, r in enumerate(tidak_acc, 1):
    vals = [r['no'], r['nama'], r['lp'], r['ktp'], r['kk'], r['alamat'], r['alasan']]
    for j, v in enumerate(vals):
        a = 'left' if j in [1, 5, 6] else 'center'
        bg = WHITE if i%2==1 else SURFACE_BG
        shade(table.rows[i].cells[j], bg)
        txt(table.rows[i].cells[j], v, size=8, align=a, bold=(j==6))

doc.save(output_path)
print(f"Word generated: {output_path}")
